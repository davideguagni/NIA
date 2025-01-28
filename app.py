from flask import Flask, request, jsonify
from docx import Document
import speech_recognition as sr

app = Flask(__name__)

# Lista per gli appuntamenti
appointments = []

# Home route
@app.route('/')
def home():
    return "Assistente Virtuale Attivo!"

# Rotta per la chat
@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json.get('message', '')
    response = f"Tu hai detto: {user_input}. Io sono qui per aiutarti!"
    return jsonify({"response": response})

# Rotta per il riconoscimento vocale
@app.route('/speech', methods=['POST'])
def speech_to_text():
    recognizer = sr.Recognizer()
    # Assumi che il file audio sia caricato come un file multipart/form-data
    audio_file = request.files['audio']
    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio)
        return jsonify({"text": text})
    except sr.UnknownValueError:
        return jsonify({"error": "Non sono riuscito a capire il comando"}), 400
    except sr.RequestError:
        return jsonify({"error": "Errore nella richiesta al servizio di riconoscimento"}), 500

# Rotta per aggiungere un appuntamento
@app.route('/appointments', methods=['POST'])
def add_appointment():
    appointment = request.json.get('appointment')
    if not appointment:
        return jsonify({"error": "Appuntamento non specificato"}), 400
    appointments.append(appointment)
    return jsonify({"message": "Appuntamento aggiunto!"}), 201

# Rotta per visualizzare tutti gli appuntamenti
@app.route('/appointments', methods=['GET'])
def get_appointments():
    return jsonify({"appointments": appointments})

# Rotta per generare il report
@app.route('/generate_report', methods=['POST'])
def generate_report():
    data = request.json.get('data')
    if not data:
        return jsonify({"error": "Nessun dato fornito per il report"}), 400
    doc = Document()
    doc.add_heading('Report', 0)
    doc.add_paragraph(data)
    doc.save('report.docx')
    return jsonify({"message": "Report generato!"}), 200

if __name__ == '__main__':
    # Avvia il server Flask
    app.run(host='0.0.0.0', port=5000)

from flask import Flask, request, jsonify
from assistant import get_assistant_response  # Questo è il nostro modulo dell'assistente
import speech_recognition as sr

app = Flask(__name__)

appointments = []  # Lista per gli appuntamenti

@app.route('/')
def home():
    return "Assistente Virtuale Attivo!"

# Rotta per la chat
@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json.get('message', '')
    
    # Richiama l'assistente per ottenere una risposta
    assistant_response = get_assistant_response(user_input)
    
    return jsonify({"response": assistant_response})

# Rotta per il riconoscimento vocale
@app.route('/speech', methods=['POST'])
def speech_to_text():
    recognizer = sr.Recognizer()
    # Assumi che il file audio sia caricato come un file multipart/form-data
    audio_file = request.files['audio']
    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio)
        # Richiama l'assistente con il testo riconosciuto
        assistant_response = get_assistant_response(text)
        return jsonify({"text": text, "response": assistant_response})
    except sr.UnknownValueError:
        return jsonify({"error": "Non sono riuscito a capire il comando"}), 400
    except sr.RequestError:
        return jsonify({"error": "Errore nella richiesta al servizio di riconoscimento"}), 500

# Rotta per aggiungere un appuntamento
@app.route('/appointments', methods=['POST'])
def add_appointment():
    appointment = request.json.get('appointment')
    if not appointment:
        return jsonify({"error": "Appuntamento non specificato"}), 400
    appointments.append(appointment)
    return jsonify({"message": "Appuntamento aggiunto!"}), 201

# Rotta per visualizzare tutti gli appuntamenti
@app.route('/appointments', methods=['GET'])
def get_appointments():
    return jsonify({"appointments": appointments})

# Rotta per generare il report
@app.route('/generate_report', methods=['POST'])
def generate_report():
    data = request.json.get('data')
    if not data:
        return jsonify({"error": "Nessun dato fornito per il report"}), 400
    doc = Document()
    doc.add_heading('Report', 0)
    doc.add_paragraph(data)
    doc.save('report.docx')
    return jsonify({"message": "Report generato!"}), 200

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)






